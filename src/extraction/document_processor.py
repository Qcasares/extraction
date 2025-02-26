"""Document processor module."""

import json
import logging
import os
import re
import xml.etree.ElementTree as ET
from csv import DictWriter
from dataclasses import asdict, dataclass, field
from datetime import datetime
from enum import Enum
from functools import lru_cache
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Set, Tuple, Union

import docx
import requests
import spacy
from docx.document import Document as DocxDocument
from tqdm import tqdm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""

    title: str = ""
    author: str = ""
    created_date: str = ""
    last_modified_date: str = ""
    file_path: str = ""
    filename: str = ""


@dataclass
class DocumentSection:
    """A section of a document with heading and content."""

    title: str
    level: int
    content: str
    entities: List[Dict[str, Any]] = field(default_factory=list)


class AIProvider(Enum):
    """Supported AI providers for document analysis."""
    
    OLLAMA = "ollama"
    OPENAI = "openai"
    ANTHROPIC = "anthropic"
    NONE = "none"
    
    @classmethod
    def from_string(cls, value: str) -> "AIProvider":
        """Convert string to AIProvider enum.
        
        Args:
            value: String representation of AI provider
            
        Returns:
            Corresponding AIProvider enum value
        """
        try:
            return cls(value.lower())
        except ValueError:
            logger.warning(f"Unknown AI provider: {value}, defaulting to OLLAMA")
            return cls.OLLAMA


@dataclass
class AIFeatures:
    """Available AI analysis features."""
    
    summary: bool = True
    topics: bool = True
    categories: bool = True
    sentiment: bool = True
    relationships: bool = True
    quality: bool = True
    suggestions: bool = True
    themes: bool = True
    insights: bool = True
    
    @classmethod
    def from_string(cls, features_str: str) -> "AIFeatures":
        """Parse comma-separated AI features string.
        
        Args:
            features_str: Comma-separated list of AI features or "all"/"none"
            
        Returns:
            AIFeatures instance with appropriate flags set
        """
        if features_str.lower() == "all":
            return cls()
        
        if features_str.lower() == "none":
            return cls(
                summary=False,
                topics=False,
                categories=False,
                sentiment=False,
                relationships=False,
                quality=False,
                suggestions=False,
                themes=False,
                insights=False
            )
            
        # Start with all features disabled
        features = cls(
            summary=False,
            topics=False,
            categories=False,
            sentiment=False,
            relationships=False,
            quality=False,
            suggestions=False,
            themes=False,
            insights=False
        )
        
        # Enable only specified features
        for feature in features_str.split(","):
            feature = feature.strip().lower()
            if hasattr(features, feature):
                setattr(features, feature, True)
                
        return features
    
    def as_dict(self) -> Dict[str, bool]:
        """Convert to dictionary.
        
        Returns:
            Dictionary mapping feature names to boolean values
        """
        return asdict(self)
    
    def is_any_enabled(self) -> bool:
        """Check if any feature is enabled.
        
        Returns:
            True if at least one feature is enabled, False otherwise
        """
        return any(asdict(self).values())


@dataclass
class AIConfig:
    """Configuration for AI-powered document analysis."""
    
    enabled: bool = False
    provider: AIProvider = AIProvider.OLLAMA
    model: str = "llama3"
    temperature: float = 0.1
    max_tokens: int = 2000
    timeout: int = 60
    api_base: str = "http://localhost:11434/api"
    api_key: str = ""
    features: AIFeatures = field(default_factory=AIFeatures)
    cache_enabled: bool = True
    cache_size: int = 100  # Number of responses to cache
    
    def is_enabled(self) -> bool:
        """Check if AI analysis is enabled and configured.
        
        Returns:
            True if AI analysis is enabled and at least one feature is enabled
        """
        return self.enabled and self.provider != AIProvider.NONE and self.features.is_any_enabled()


@dataclass
class AiAnalysis:
    """AI-generated analysis of a document."""
    
    summary: str = ""
    key_topics: List[str] = field(default_factory=list)
    content_categories: List[str] = field(default_factory=list)
    sentiment: str = ""
    entity_relationships: List[Dict[str, Any]] = field(default_factory=list)
    document_quality_score: float = 0.0
    improvement_suggestions: List[str] = field(default_factory=list)
    themes: List[Dict[str, Any]] = field(default_factory=list)
    key_insights: List[str] = field(default_factory=list)
    provider: AIProvider = AIProvider.NONE
    model: str = ""
    processing_time: float = 0.0


@dataclass
class ProcessedDocument:
    """Document with extracted and processed content."""

    metadata: DocumentMetadata
    sections: List[DocumentSection] = field(default_factory=list)
    raw_text: str = ""
    ai_analysis: Optional[AiAnalysis] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        result = {
            "metadata": asdict(self.metadata),
            "sections": [asdict(section) for section in self.sections],
            "raw_text": self.raw_text,
        }
        
        if self.ai_analysis:
            result["ai_analysis"] = asdict(self.ai_analysis)
            
        return result


class DocumentProcessor:
    """Processor for Word documents."""

    def __init__(
        self,
        input_directory: str,
        output_format: str = "json",
        use_ai: bool = False,
        ai_provider: str = "ollama",
        ai_model: Optional[str] = None,
        api_base: Optional[str] = None,
        api_key: str = "",
        temperature: float = 0.1,
        max_tokens: int = 2000,
        timeout: int = 60,
        ai_features: str = "all",
        ai_cache_enabled: bool = True,
        ai_cache_size: int = 100,
    ) -> None:
        """Initialize the document processor.

        Args:
            input_directory: Directory to process documents from
            output_format: Output format (json, xml, csv)
            use_ai: Whether to use AI for analysis
            ai_provider: AI provider to use (ollama, openai, anthropic, none)
            ai_model: Model to use (provider-specific)
            api_base: Base URL for API calls
            api_key: API key for authentication (needed for OpenAI and Anthropic)
            temperature: Temperature parameter for AI generation (0.0-1.0)
            max_tokens: Maximum number of tokens for AI response
            timeout: Timeout in seconds for API calls
            ai_features: Comma-separated list of AI analysis features to enable
            ai_cache_enabled: Whether to cache AI responses
            ai_cache_size: Number of AI responses to cache
        """
        self.input_directory = Path(input_directory)
        self.output_format = output_format.lower()
        self.nlp = spacy.load("en_core_web_sm")
        self.processed_documents: List[ProcessedDocument] = []

        # Configure AI
        provider = AIProvider.from_string(ai_provider)
        features = AIFeatures.from_string(ai_features)
        
        # Set provider-specific defaults if not provided
        if api_base is None:
            if provider == AIProvider.OLLAMA:
                api_base = "http://localhost:11434/api"
            elif provider == AIProvider.OPENAI:
                api_base = "https://api.openai.com/v1"
            elif provider == AIProvider.ANTHROPIC:
                api_base = "https://api.anthropic.com/v1"
            else:
                api_base = ""
                
        # Set provider-specific default models if not provided
        if ai_model is None:
            if provider == AIProvider.OLLAMA:
                ai_model = "llama3"
            elif provider == AIProvider.OPENAI:
                ai_model = "gpt-4o"
            elif provider == AIProvider.ANTHROPIC:
                ai_model = "claude-3-opus-20240229"
            else:
                ai_model = ""
        
        self.ai_config = AIConfig(
            enabled=use_ai,
            provider=provider,
            model=ai_model,
            temperature=temperature,
            max_tokens=max_tokens,
            timeout=timeout,
            api_base=api_base,
            api_key=api_key,
            features=features,
            cache_enabled=ai_cache_enabled,
            cache_size=ai_cache_size,
        )

        # Validate input directory
        if not self.input_directory.exists():
            raise ValueError(f"Input directory does not exist: {input_directory}")
        if not self.input_directory.is_dir():
            raise ValueError(f"Input path is not a directory: {input_directory}")

        # Validate output format
        if self.output_format not in ("json", "xml", "csv"):
            raise ValueError(
                f"Invalid output format: {output_format}. Must be json, xml, or csv."
            )

        logger.info(
            "Initialized DocumentProcessor with input_directory=%s, output_format=%s, ai_enabled=%s, "
            "provider=%s, model=%s, features=%s",
            input_directory,
            output_format,
            self.ai_config.enabled,
            self.ai_config.provider.value,
            self.ai_config.model,
            ",".join(k for k, v in self.ai_config.features.as_dict().items() if v),
        )

    def find_documents(self) -> List[Path]:
        """Find all Word documents in the input directory.

        Returns:
            List of paths to Word documents
        """
        logger.info("Searching for .docx files in %s", self.input_directory)
        docx_files: List[Path] = []

        for path in self.input_directory.glob("**/*.docx"):
            if path.is_file():
                docx_files.append(path)

        logger.info("Found %d .docx files", len(docx_files))
        return docx_files

    def extract_metadata(self, doc: DocxDocument, file_path: Path) -> DocumentMetadata:
        """Extract metadata from a Word document.

        Args:
            doc: Word document
            file_path: Path to the document

        Returns:
            Document metadata
        """
        properties = doc.core_properties
        metadata = DocumentMetadata(
            title=properties.title if properties.title else "",
            author=properties.author if properties.author else "",
            created_date=properties.created.isoformat() if properties.created else "",
            last_modified_date=(
                properties.modified.isoformat() if properties.modified else ""
            ),
            file_path=str(file_path.absolute()),
            filename=file_path.name,
        )
        return metadata

    def extract_text(self, doc: DocxDocument) -> str:
        """Extract text from a Word document.

        Args:
            doc: Word document

        Returns:
            Extracted text
        """
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    def identify_sections(self, doc: DocxDocument) -> List[DocumentSection]:
        """Identify document sections based on headings.

        Args:
            doc: Word document

        Returns:
            List of document sections
        """
        sections: List[DocumentSection] = []
        current_section: Optional[DocumentSection] = None
        content_buffer: List[str] = []

        for paragraph in doc.paragraphs:
            # Check if paragraph is a heading
            # Check if paragraph has a style and if it's a heading
            if paragraph.style and hasattr(paragraph.style, 'name') and paragraph.style.name and paragraph.style.name.startswith("Heading"):
                # Extract heading level (Heading 1, Heading 2, etc.)
                try:
                    level = int(paragraph.style.name.split()[-1])
                except ValueError:
                    level = 0

                # If we have a previous section, save its content
                if current_section is not None:
                    current_section.content = "\n".join(content_buffer)
                    sections.append(current_section)

                # Create a new section
                current_section = DocumentSection(
                    title=paragraph.text,
                    level=level,
                    content="",
                )
                content_buffer = []
            elif current_section is not None:
                # Add paragraph to current section
                if paragraph.text.strip():
                    content_buffer.append(paragraph.text)
            else:
                # Text before any heading - create a default section
                if paragraph.text.strip() and not current_section:
                    current_section = DocumentSection(
                        title="Introduction",
                        level=0,
                        content="",
                    )
                    content_buffer.append(paragraph.text)

        # Don't forget the last section
        if current_section is not None:
            current_section.content = "\n".join(content_buffer)
            sections.append(current_section)

        return sections

    def analyze_content(self, processed_doc: ProcessedDocument) -> None:
        """Analyze document content to extract entities and structure.

        Args:
            processed_doc: Processed document to analyze
        """
        # Process each section with spaCy
        for section in processed_doc.sections:
            if not section.content:
                continue

            doc = self.nlp(section.content)

            # Extract named entities
            entities = []
            for ent in doc.ents:
                entities.append(
                    {
                        "text": ent.text,
                        "label": ent.label_,
                        "start": ent.start_char,
                        "end": ent.end_char,
                    }
                )

            section.entities = entities

    def _generate_prompt_template(self, processed_doc: ProcessedDocument) -> str:
        """Generate a prompt template for AI analysis.
        
        Args:
            processed_doc: The document to analyze
            
        Returns:
            A formatted prompt template
        """
        # Extract section titles for context
        section_titles = [section.title for section in processed_doc.sections]
        section_titles_str = "\n".join([f"- {title}" for title in section_titles])
        
        # Extract entities for context
        all_entities = []
        for section in processed_doc.sections:
            all_entities.extend([entity["text"] for entity in section.entities])
        unique_entities = list(set(all_entities))[:20]  # Limit to avoid token limits
        entities_str = ", ".join(unique_entities)
        
        # Get document content (truncated if necessary)
        max_content_length = 3000  # Adjust based on provider token limits
        doc_content = processed_doc.raw_text[:max_content_length]
        if len(processed_doc.raw_text) > max_content_length:
            doc_content += "..."
        
        # Build JSON structure based on enabled features
        json_structure_parts = []
        features = self.ai_config.features
        
        if features.summary:
            json_structure_parts.append('"summary": "A concise 2-3 sentence summary of the entire document"')
            
        if features.topics:
            json_structure_parts.append('"key_topics": ["topic1", "topic2", "topic3", ...]')
            
        if features.categories:
            json_structure_parts.append('"content_categories": ["category1", "category2", ...]')
            
        if features.sentiment:
            json_structure_parts.append('"sentiment": "The overall sentiment of the document (positive, negative, neutral, or mixed)"')
            
        if features.relationships:
            json_structure_parts.append(
                '"entity_relationships": [\n'
                '    {\n'
                '        "entity1": "Name of first entity",\n'
                '        "entity2": "Name of second entity",\n'
                '        "relationship": "Description of relationship"\n'
                '    },\n'
                '    ...\n'
                ']'
            )
            
        if features.quality:
            json_structure_parts.append('"document_quality_score": A score from 0.0 to 10.0 rating the quality of the document')
            
        if features.suggestions:
            json_structure_parts.append(
                '"improvement_suggestions": [\n'
                '    "Suggestion 1",\n'
                '    "Suggestion 2",\n'
                '    ...\n'
                ']'
            )
            
        if features.themes:
            json_structure_parts.append(
                '"themes": [\n'
                '    {\n'
                '        "theme": "Theme name",\n'
                '        "relevance_score": Score from 0.0 to 1.0,\n'
                '        "sections": ["Section where theme appears", ...]\n'
                '    },\n'
                '    ...\n'
                ']'
            )
            
        if features.insights:
            json_structure_parts.append(
                '"key_insights": [\n'
                '    "Key insight 1",\n'
                '    "Key insight 2",\n'
                '    ...\n'
                ']'
            )
            
        # Join all parts with commas
        json_structure = ",\n".join(json_structure_parts)
        
        # Prepare the comprehensive analysis prompt
        return f"""
        Please perform a comprehensive analysis of this document. Return your response in JSON format according to the structure below.
        
        Document Title: {processed_doc.metadata.title}
        Author: {processed_doc.metadata.author}
        Created Date: {processed_doc.metadata.created_date}
        
        Document Content (partial):
        {doc_content}
        
        Document Sections:
        {section_titles_str}
        
        Key Entities Identified:
        {entities_str}
        
        Perform a comprehensive analysis and return a JSON response with the following structure:
        {{
{json_structure}
        }}
        
        IMPORTANT: Respond ONLY with the JSON. Do not include any other text.
        """
        
    @lru_cache(maxsize=100)
    def _get_ai_response_cached(
        self, 
        prompt: str, 
        provider: AIProvider, 
        model: str, 
        temperature: float, 
        max_tokens: int
    ) -> Tuple[bool, str]:
        """Get AI response with caching to avoid redundant calls.
        
        Args:
            prompt: The prompt to send to the AI
            provider: AI provider to use
            model: Model to use
            temperature: Temperature parameter
            max_tokens: Maximum tokens to generate
            
        Returns:
            Tuple of (success, response)
        """
        try:
            if provider == AIProvider.OLLAMA:
                return self._get_ollama_response(prompt, model, temperature, max_tokens)
            elif provider == AIProvider.OPENAI:
                return self._get_openai_response(prompt, model, temperature, max_tokens)
            elif provider == AIProvider.ANTHROPIC:
                return self._get_anthropic_response(prompt, model, temperature, max_tokens)
            else:
                return False, "No AI provider configured"
        except Exception as e:
            logger.error(f"Error getting AI response: {str(e)}")
            return False, f"Error: {str(e)}"
    
    def _get_ollama_response(
        self, prompt: str, model: str, temperature: float, max_tokens: int
    ) -> Tuple[bool, str]:
        """Get response from Ollama API.
        
        Args:
            prompt: The prompt to send
            model: Ollama model to use
            temperature: Temperature parameter
            max_tokens: Maximum tokens to generate
            
        Returns:
            Tuple of (success, response)
        """
        try:
            response = requests.post(
                f"{self.ai_config.api_base}/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "temperature": temperature,
                    "top_p": 0.9,
                    "max_tokens": max_tokens,
                },
                timeout=self.ai_config.timeout,
            )
            
            if response.status_code == 200:
                result = response.json()
                return True, result.get("response", "")
            else:
                logger.warning(
                    f"Ollama API error: {response.status_code}, {response.text}"
                )
                return False, f"API error: {response.status_code}"
        except Exception as e:
            logger.error(f"Ollama API exception: {str(e)}")
            return False, f"Error: {str(e)}"
    
    def _get_openai_response(
        self, prompt: str, model: str, temperature: float, max_tokens: int
    ) -> Tuple[bool, str]:
        """Get response from OpenAI API.
        
        Args:
            prompt: The prompt to send
            model: OpenAI model to use
            temperature: Temperature parameter
            max_tokens: Maximum tokens to generate
            
        Returns:
            Tuple of (success, response)
        """
        try:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.ai_config.api_key}"
            }
            
            response = requests.post(
                f"{self.ai_config.api_base}/chat/completions",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": "You are a document analysis assistant that responds in JSON format."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": temperature,
                    "max_tokens": max_tokens,
                    "response_format": {"type": "json_object"}
                },
                headers=headers,
                timeout=self.ai_config.timeout,
            )
            
            if response.status_code == 200:
                result = response.json()
                return True, result.get("choices", [{}])[0].get("message", {}).get("content", "")
            else:
                logger.warning(
                    f"OpenAI API error: {response.status_code}, {response.text}"
                )
                return False, f"API error: {response.status_code}"
        except Exception as e:
            logger.error(f"OpenAI API exception: {str(e)}")
            return False, f"Error: {str(e)}"
    
    def _get_anthropic_response(
        self, prompt: str, model: str, temperature: float, max_tokens: int
    ) -> Tuple[bool, str]:
        """Get response from Anthropic API.
        
        Args:
            prompt: The prompt to send
            model: Anthropic model to use
            temperature: Temperature parameter
            max_tokens: Maximum tokens to generate
            
        Returns:
            Tuple of (success, response)
        """
        try:
            headers = {
                "Content-Type": "application/json",
                "x-api-key": self.ai_config.api_key,
                "anthropic-version": "2023-06-01"
            }
            
            response = requests.post(
                f"{self.ai_config.api_base}/messages",
                json={
                    "model": model,
                    "system": "You are a document analysis assistant that responds in JSON format.",
                    "messages": [
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": temperature,
                    "max_tokens": max_tokens
                },
                headers=headers,
                timeout=self.ai_config.timeout,
            )
            
            if response.status_code == 200:
                result = response.json()
                return True, result.get("content", [{}])[0].get("text", "")
            else:
                logger.warning(
                    f"Anthropic API error: {response.status_code}, {response.text}"
                )
                return False, f"API error: {response.status_code}"
        except Exception as e:
            logger.error(f"Anthropic API exception: {str(e)}")
            return False, f"Error: {str(e)}"
            
    def _parse_ai_response(self, response_text: str) -> Dict[str, Any]:
        """Parse AI response text to extract JSON data.
        
        Args:
            response_text: Raw response text from AI
            
        Returns:
            Parsed JSON data or empty dict on failure
        """
        try:
            # Try direct JSON parsing first
            return json.loads(response_text)
        except json.JSONDecodeError:
            # Find JSON in response (in case model outputs additional text)
            json_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
            matches = re.findall(json_pattern, response_text)
            
            if matches:
                # Use the largest match as it's most likely to be our complete response
                try:
                    json_str = max(matches, key=len)
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    logger.warning("Failed to parse JSON from largest match")
            
            # Try a more aggressive approach to find valid JSON
            for potential_json in matches:
                try:
                    return json.loads(potential_json)
                except json.JSONDecodeError:
                    continue
                    
            # If we get here, no valid JSON was found
            logger.warning("No valid JSON found in response")
            return {}

    def analyze_with_ai(self, processed_doc: ProcessedDocument) -> None:
        """Use AI to enhance document analysis with advanced features.

        Args:
            processed_doc: Processed document to analyze
        """
        # Skip if AI is not configured or no features are enabled
        if not self.ai_config.is_enabled():
            return

        import time
        start_time = time.time()
        
        try:
            # Generate prompt
            prompt = self._generate_prompt_template(processed_doc)
            
            # Get AI response (with caching if enabled)
            if self.ai_config.cache_enabled:
                success, ai_response = self._get_ai_response_cached(
                    prompt=prompt,
                    provider=self.ai_config.provider,
                    model=self.ai_config.model,
                    temperature=self.ai_config.temperature,
                    max_tokens=self.ai_config.max_tokens
                )
            else:
                # Get response without caching
                if self.ai_config.provider == AIProvider.OLLAMA:
                    success, ai_response = self._get_ollama_response(
                        prompt, self.ai_config.model, self.ai_config.temperature, self.ai_config.max_tokens
                    )
                elif self.ai_config.provider == AIProvider.OPENAI:
                    success, ai_response = self._get_openai_response(
                        prompt, self.ai_config.model, self.ai_config.temperature, self.ai_config.max_tokens
                    )
                elif self.ai_config.provider == AIProvider.ANTHROPIC:
                    success, ai_response = self._get_anthropic_response(
                        prompt, self.ai_config.model, self.ai_config.temperature, self.ai_config.max_tokens
                    )
                else:
                    success = False
                    ai_response = "No AI provider configured"
            
            if success:
                # Parse JSON response
                json_data = self._parse_ai_response(ai_response)
                
                if json_data:
                    # Create AiAnalysis object from JSON response
                    ai_analysis = AiAnalysis(
                        summary=json_data.get("summary", ""),
                        key_topics=json_data.get("key_topics", []),
                        content_categories=json_data.get("content_categories", []),
                        sentiment=json_data.get("sentiment", ""),
                        entity_relationships=json_data.get("entity_relationships", []),
                        document_quality_score=json_data.get("document_quality_score", 0.0),
                        improvement_suggestions=json_data.get("improvement_suggestions", []),
                        themes=json_data.get("themes", []),
                        key_insights=json_data.get("key_insights", []),
                        provider=self.ai_config.provider,
                        model=self.ai_config.model,
                        processing_time=time.time() - start_time
                    )
                    
                    # Attach AI analysis to processed document
                    processed_doc.ai_analysis = ai_analysis
                    
                    logger.info(
                        "AI analysis completed for %s using %s model %s (%.2fs)",
                        processed_doc.metadata.filename,
                        self.ai_config.provider.value,
                        self.ai_config.model,
                        ai_analysis.processing_time
                    )
                else:
                    # Try simple extraction as fallback
                    logger.warning("Failed to parse JSON response, trying simple extraction")
                    self._extract_simple_analysis(processed_doc, ai_response)
            else:
                logger.warning(f"Failed to get AI analysis: {ai_response}")
                
        except Exception as e:
            logger.error(f"Error during AI analysis: {str(e)}")
            # Log exception details for debugging
            import traceback
            logger.debug(f"AI analysis exception details: {traceback.format_exc()}")
            # Fall back to rule-based analysis
            logger.info("Falling back to rule-based analysis")
            
    def _extract_simple_analysis(self, processed_doc: ProcessedDocument, ai_response: str) -> None:
        """Extract simple analysis from AI response when JSON parsing fails.
        
        Args:
            processed_doc: The document being processed
            ai_response: The raw AI response text
        """
        # Create a basic AI analysis with whatever we can extract
        ai_analysis = AiAnalysis()
        
        # Try to extract summary (look for "summary" section)
        summary_match = re.search(r"summary[\"']?\s*:?\s*[\"']([^\"']+)[\"']", ai_response, re.IGNORECASE)
        if summary_match:
            ai_analysis.summary = summary_match.group(1).strip()
            
        # Try to extract key topics
        topics_match = re.search(r"key[_\s]topics[\"']?\s*:?\s*\[(.*?)\]", ai_response, re.IGNORECASE | re.DOTALL)
        if topics_match:
            topics_text = topics_match.group(1)
            # Extract quoted strings
            topics = re.findall(r"[\"']([^\"']+)[\"']", topics_text)
            ai_analysis.key_topics = topics
            
        # Try to extract sentiment
        sentiment_match = re.search(r"sentiment[\"']?\s*:?\s*[\"']([^\"']+)[\"']", ai_response, re.IGNORECASE)
        if sentiment_match:
            ai_analysis.sentiment = sentiment_match.group(1).strip()
            
        # Add the basic analysis to the document
        processed_doc.ai_analysis = ai_analysis
        logger.info("Simple AI analysis extracted for %s", processed_doc.metadata.filename)

    def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process a single document.

        Args:
            file_path: Path to the document

        Returns:
            Processed document
        """
        logger.info("Processing document: %s", file_path)

        try:
            doc = docx.Document(str(file_path))

            # Extract metadata
            metadata = self.extract_metadata(doc, file_path)

            # Extract raw text
            raw_text = self.extract_text(doc)

            # Identify sections
            sections = self.identify_sections(doc)

            # Create processed document
            processed_doc = ProcessedDocument(
                metadata=metadata,
                sections=sections,
                raw_text=raw_text,
            )

            # Analyze content with spaCy
            self.analyze_content(processed_doc)

            # Use AI for enhanced analysis if enabled
            if self.ai_config.is_enabled():
                self.analyze_with_ai(processed_doc)

            return processed_doc

        except Exception as e:
            logger.error("Error processing document %s: %s", file_path, str(e))
            import traceback
            logger.debug(f"Document processing exception details: {traceback.format_exc()}")
            
            # Return an empty document with error information
            return ProcessedDocument(
                metadata=DocumentMetadata(
                    file_path=str(file_path),
                    filename=file_path.name,
                )
            )

    def process_all(self) -> List[ProcessedDocument]:
        """Process all documents in the input directory.

        Returns:
            List of processed documents
        """
        documents = self.find_documents()
        self.processed_documents = []

        for file_path in tqdm(documents, desc="Processing documents"):
            processed_doc = self.process_document(file_path)
            self.processed_documents.append(processed_doc)

        return self.processed_documents

    def save_results(self, output_path: str) -> None:
        """Save processed results to the specified path.

        Args:
            output_path: Path to save results to
        """
        output_dir = Path(output_path)
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if self.output_format == "json":
            self._save_as_json(output_dir, timestamp)
        elif self.output_format == "xml":
            self._save_as_xml(output_dir, timestamp)
        elif self.output_format == "csv":
            self._save_as_csv(output_dir, timestamp)

        logger.info(
            "Saved %d documents to %s in %s format",
            len(self.processed_documents),
            output_dir,
            self.output_format,
        )

    def _save_as_json(self, output_dir: Path, timestamp: str) -> None:
        """Save results as JSON.

        Args:
            output_dir: Directory to save to
            timestamp: Timestamp for filename
        """
        # Save individual documents
        for doc in self.processed_documents:
            filename = f"{doc.metadata.filename.replace('.docx', '')}_{timestamp}.json"
            with open(output_dir / filename, "w") as f:
                json.dump(doc.to_dict(), f, indent=2)

        # Save summary file with all documents
        summary_file = output_dir / f"document_summary_{timestamp}.json"
        with open(summary_file, "w") as f:
            json.dump([doc.to_dict() for doc in self.processed_documents], f, indent=2)

    def _save_as_xml(self, output_dir: Path, timestamp: str) -> None:
        """Save results as XML.

        Args:
            output_dir: Directory to save to
            timestamp: Timestamp for filename
        """
        for doc in self.processed_documents:
            filename = f"{doc.metadata.filename.replace('.docx', '')}_{timestamp}.xml"

            # Create XML structure
            root = ET.Element("document")

            # Add metadata
            metadata_elem = ET.SubElement(root, "metadata")
            for key, value in asdict(doc.metadata).items():
                ET.SubElement(metadata_elem, key).text = str(value)

            # Add sections
            sections_elem = ET.SubElement(root, "sections")
            for section in doc.sections:
                section_elem = ET.SubElement(sections_elem, "section")
                ET.SubElement(section_elem, "title").text = section.title
                ET.SubElement(section_elem, "level").text = str(section.level)
                ET.SubElement(section_elem, "content").text = section.content

                # Add entities
                entities_elem = ET.SubElement(section_elem, "entities")
                for entity in section.entities:
                    entity_elem = ET.SubElement(entities_elem, "entity")
                    for key, value in entity.items():
                        ET.SubElement(entity_elem, key).text = str(value)
                        
            # Add AI analysis if available
            if doc.ai_analysis:
                ai_analysis_elem = ET.SubElement(root, "ai_analysis")
                
                # Add simple fields
                ET.SubElement(ai_analysis_elem, "summary").text = doc.ai_analysis.summary
                ET.SubElement(ai_analysis_elem, "sentiment").text = doc.ai_analysis.sentiment
                ET.SubElement(ai_analysis_elem, "document_quality_score").text = str(doc.ai_analysis.document_quality_score)
                
                # Add list fields
                topics_elem = ET.SubElement(ai_analysis_elem, "key_topics")
                for topic in doc.ai_analysis.key_topics:
                    ET.SubElement(topics_elem, "topic").text = topic
                    
                categories_elem = ET.SubElement(ai_analysis_elem, "content_categories")
                for category in doc.ai_analysis.content_categories:
                    ET.SubElement(categories_elem, "category").text = category
                    
                suggestions_elem = ET.SubElement(ai_analysis_elem, "improvement_suggestions")
                for suggestion in doc.ai_analysis.improvement_suggestions:
                    ET.SubElement(suggestions_elem, "suggestion").text = suggestion
                    
                insights_elem = ET.SubElement(ai_analysis_elem, "key_insights")
                for insight in doc.ai_analysis.key_insights:
                    ET.SubElement(insights_elem, "insight").text = insight
                
                # Add complex fields
                relationships_elem = ET.SubElement(ai_analysis_elem, "entity_relationships")
                for relationship in doc.ai_analysis.entity_relationships:
                    relationship_elem = ET.SubElement(relationships_elem, "relationship")
                    for key, value in relationship.items():
                        ET.SubElement(relationship_elem, key).text = str(value)
                        
                themes_elem = ET.SubElement(ai_analysis_elem, "themes")
                for theme in doc.ai_analysis.themes:
                    theme_elem = ET.SubElement(themes_elem, "theme")
                    for key, value in theme.items():
                        if key == "sections" and isinstance(value, list):
                            sections_item = ET.SubElement(theme_elem, "sections")
                            for section in value:
                                ET.SubElement(sections_item, "section").text = str(section)
                        else:
                            ET.SubElement(theme_elem, key).text = str(value)

            # Write to file
            tree = ET.ElementTree(root)
            # Only use indent if available (Python 3.9+)
            if hasattr(ET, "indent"):
                ET.indent(tree, space="  ")
            tree.write(output_dir / filename, encoding="utf-8", xml_declaration=True)

    def _save_as_csv(self, output_dir: Path, timestamp: str) -> None:
        """Save results as CSV.

        Args:
            output_dir: Directory to save to
            timestamp: Timestamp for filename
        """
        # CSV for document metadata
        metadata_file = output_dir / f"document_metadata_{timestamp}.csv"
        with open(metadata_file, "w", newline="") as f:
            fieldnames = [
                "filename",
                "title",
                "author",
                "created_date",
                "last_modified_date",
                "file_path",
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for doc in self.processed_documents:
                writer.writerow(asdict(doc.metadata))

        # CSV for sections
        sections_file = output_dir / f"document_sections_{timestamp}.csv"
        with open(sections_file, "w", newline="") as f:
            fieldnames = ["document_filename", "section_title", "level", "content"]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for doc in self.processed_documents:
                for section in doc.sections:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "section_title": section.title,
                            "level": section.level,
                            "content": section.content,
                        }
                    )

        # CSV for entities
        entities_file = output_dir / f"document_entities_{timestamp}.csv"
        with open(entities_file, "w", newline="") as f:
            fieldnames = [
                "document_filename",
                "section_title",
                "entity_text",
                "entity_label",
                "entity_start",
                "entity_end",
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for doc in self.processed_documents:
                for section in doc.sections:
                    for entity in section.entities:
                        writer.writerow(
                            {
                                "document_filename": doc.metadata.filename,
                                "section_title": section.title,
                                "entity_text": entity["text"],
                                "entity_label": entity["label"],
                                "entity_start": entity["start"],
                                "entity_end": entity["end"],
                            }
                        )
                        
        # CSV for AI analysis summary
        ai_analysis_file = output_dir / f"document_ai_analysis_{timestamp}.csv"
        with open(ai_analysis_file, "w", newline="") as f:
            fieldnames = [
                "document_filename",
                "summary", 
                "sentiment",
                "document_quality_score",
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for doc in self.processed_documents:
                if doc.ai_analysis:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "summary": doc.ai_analysis.summary,
                            "sentiment": doc.ai_analysis.sentiment,
                            "document_quality_score": doc.ai_analysis.document_quality_score,
                        }
                    )
        
        # CSV for AI analysis topics and categories
        topics_file = output_dir / f"document_ai_topics_{timestamp}.csv"
        with open(topics_file, "w", newline="") as f:
            fieldnames = [
                "document_filename",
                "type",  # "topic", "category", "insight", "suggestion"
                "value", 
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            for doc in self.processed_documents:
                if not doc.ai_analysis:
                    continue
                    
                # Write topics
                for topic in doc.ai_analysis.key_topics:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "type": "topic",
                            "value": topic,
                        }
                    )
                
                # Write categories
                for category in doc.ai_analysis.content_categories:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "type": "category",
                            "value": category,
                        }
                    )
                
                # Write insights
                for insight in doc.ai_analysis.key_insights:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "type": "insight",
                            "value": insight,
                        }
                    )
                
                # Write suggestions
                for suggestion in doc.ai_analysis.improvement_suggestions:
                    writer.writerow(
                        {
                            "document_filename": doc.metadata.filename,
                            "type": "suggestion",
                            "value": suggestion,
                        }
                    )
                    
        # CSV for entity relationships
        relationships_file = output_dir / f"document_ai_relationships_{timestamp}.csv"
        with open(relationships_file, "w", newline="") as f:
            fieldnames = [
                "document_filename",
                "entity1",
                "entity2",
                "relationship",
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            for doc in self.processed_documents:
                if not doc.ai_analysis:
                    continue
                    
                for relationship in doc.ai_analysis.entity_relationships:
                    if "entity1" in relationship and "entity2" in relationship and "relationship" in relationship:
                        writer.writerow(
                            {
                                "document_filename": doc.metadata.filename,
                                "entity1": relationship["entity1"],
                                "entity2": relationship["entity2"],
                                "relationship": relationship["relationship"],
                            }
                        )
                        
        # CSV for themes
        themes_file = output_dir / f"document_ai_themes_{timestamp}.csv"
        with open(themes_file, "w", newline="") as f:
            fieldnames = [
                "document_filename",
                "theme",
                "relevance_score",
                "sections",
            ]
            writer = DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            for doc in self.processed_documents:
                if not doc.ai_analysis:
                    continue
                    
                for theme in doc.ai_analysis.themes:
                    if "theme" in theme:
                        writer.writerow(
                            {
                                "document_filename": doc.metadata.filename,
                                "theme": theme.get("theme", ""),
                                "relevance_score": theme.get("relevance_score", 0.0),
                                "sections": ", ".join(theme.get("sections", [])),
                            }
                        )
