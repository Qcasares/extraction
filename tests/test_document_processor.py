"""Tests for document processor module."""

import os
import tempfile
from pathlib import Path
from typing import Generator
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from extraction.document_processor import (
    DocumentMetadata,
    DocumentProcessor,
    DocumentSection,
    ProcessedDocument,
)


@pytest.fixture
def sample_docx() -> Generator[Path, None, None]:
    """Create a temporary Word document for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        document = Document()
        document.add_heading("Test Document", 0)
        document.add_heading("Section 1", level=1)
        document.add_paragraph("This is content for section 1.")
        document.add_heading("Section 2", level=1)
        document.add_paragraph("This is content for section 2.")
        document.add_heading("Subsection 2.1", level=2)
        document.add_paragraph("This is content for subsection 2.1.")

        document.save(temp_file.name)
        temp_path = Path(temp_file.name)

    yield temp_path

    # Clean up
    if temp_path.exists():
        os.unlink(temp_path)


@pytest.fixture
def temp_output_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for output files."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield Path(temp_dir)


@pytest.fixture
def document_processor(sample_docx: Path) -> DocumentProcessor:
    """Create a document processor with a sample document."""
    temp_dir = sample_docx.parent
    return DocumentProcessor(input_directory=str(temp_dir), output_format="json")


def test_find_documents(document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test finding Word documents."""
    documents = document_processor.find_documents()

    # Check if our sample document is found
    filenames = [doc.name for doc in documents]
    assert sample_docx.name in filenames


def test_extract_metadata(document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test extracting metadata from a document."""
    doc = Document(str(sample_docx))
    metadata = document_processor.extract_metadata(doc, sample_docx)

    assert isinstance(metadata, DocumentMetadata)
    assert metadata.filename == sample_docx.name
    assert str(sample_docx.absolute()) in metadata.file_path


def test_extract_text(document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test extracting text from a document."""
    doc = Document(str(sample_docx))
    text = document_processor.extract_text(doc)

    assert "Test Document" in text
    assert "This is content for section 1." in text
    assert "This is content for section 2." in text
    assert "This is content for subsection 2.1." in text


def test_identify_sections(document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test identifying document sections."""
    doc = Document(str(sample_docx))
    sections = document_processor.identify_sections(doc)

    assert len(sections) == 4  # Title + 3 sections

    # Check section titles
    section_titles = [s.title for s in sections]
    assert "Section 1" in section_titles
    assert "Section 2" in section_titles
    assert "Subsection 2.1" in section_titles

    # Check section content
    for section in sections:
        if section.title == "Section 1":
            assert "This is content for section 1." in section.content
        elif section.title == "Subsection 2.1":
            assert "This is content for subsection 2.1." in section.content


def test_process_document(document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test processing a single document."""
    processed_doc = document_processor.process_document(sample_docx)

    assert isinstance(processed_doc, ProcessedDocument)
    assert processed_doc.metadata.filename == sample_docx.name
    assert len(processed_doc.sections) >= 3
    assert processed_doc.raw_text


def test_process_all(document_processor: DocumentProcessor) -> None:
    """Test processing all documents."""
    processed_docs = document_processor.process_all()

    assert len(processed_docs) >= 1
    assert isinstance(processed_docs[0], ProcessedDocument)


def test_save_results_json(document_processor: DocumentProcessor, temp_output_dir: Path) -> None:
    """Test saving results as JSON."""
    # Process a document
    document_processor.process_all()

    # Save results
    document_processor.save_results(str(temp_output_dir))

    # Check for output files
    json_files = list(temp_output_dir.glob("*.json"))
    assert len(json_files) >= 1


@patch("extraction.document_processor.requests.post")
def test_analyze_with_ai(mock_post: MagicMock, document_processor: DocumentProcessor, sample_docx: Path) -> None:
    """Test analyzing document with AI."""
    # Configure the processor to use AI
    document_processor.use_ai = True
    document_processor.ollama_model = "llama3"

    # Mock the response
    mock_response = MagicMock()
    mock_response.status_code = 200
    mock_response.json.return_value = {"response": "AI analysis output"}
    mock_post.return_value = mock_response

    # Process a document
    doc = Document(str(sample_docx))
    processed_doc = ProcessedDocument(
        metadata=document_processor.extract_metadata(doc, sample_docx),
        raw_text=document_processor.extract_text(doc),
        sections=document_processor.identify_sections(doc),
    )

    # Call the method
    document_processor.analyze_with_ai(processed_doc)

    # Verify the API was called
    mock_post.assert_called_once()


def test_save_as_xml(document_processor: DocumentProcessor, temp_output_dir: Path) -> None:
    """Test saving results as XML."""
    # Set output format to XML
    document_processor.output_format = "xml"

    # Process documents
    document_processor.process_all()

    # Save results
    document_processor.save_results(str(temp_output_dir))

    # Check for output files
    xml_files = list(temp_output_dir.glob("*.xml"))
    assert len(xml_files) >= 1


def test_save_as_csv(document_processor: DocumentProcessor, temp_output_dir: Path) -> None:
    """Test saving results as CSV."""
    # Set output format to CSV
    document_processor.output_format = "csv"

    # Process documents
    document_processor.process_all()

    # Save results
    document_processor.save_results(str(temp_output_dir))

    # Check for output files - should have at least metadata CSV
    csv_files = list(temp_output_dir.glob("*.csv"))
    assert len(csv_files) >= 1
